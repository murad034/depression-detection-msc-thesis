"""
Generate a DOCX presentation script / speaker notes for the MSc thesis defense.
One section per slide — what to say, key points, and anticipated Q&A.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────────
for sec in doc.sections:
    sec.page_width  = Cm(21.0); sec.page_height = Cm(29.7)
    sec.top_margin  = Cm(2.0);  sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5);  sec.right_margin  = Cm(2.5)

ns = doc.styles["Normal"]
ns.font.name = "Calibri"; ns.font.size = Pt(11)
ns.paragraph_format.space_before = Pt(0)
ns.paragraph_format.space_after  = Pt(4)

# ── Color helpers ─────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x3E, 0x6E)
BLUE   = RGBColor(0x2E, 0x74, 0xB5)
ORANGE = RGBColor(0xED, 0x7D, 0x31)
GREEN  = RGBColor(0x37, 0x56, 0x23)
RED    = RGBColor(0xC0, 0x00, 0x00)
GRAY   = RGBColor(0x59, 0x59, 0x59)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

def _shading(para, hex_fill):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    pPr.append(shd)

def _set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)

def _cell_txt(cell, text, bold=False, size=10, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]; p.clear(); p.alignment = align
    run = p.add_run(text)
    run.bold = bold; run.font.size = Pt(size); run.font.name = "Calibri"
    if color: run.font.color.rgb = color

def hr(doc, color="1F3E6E"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    for k,v in [("w:val","single"),("w:sz","6"),("w:space","1"),("w:color",color)]:
        bot.set(qn(k), v)
    pBdr.append(bot); pPr.append(pBdr)

def slide_banner(doc, num, title):
    """Navy-shaded slide heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(2)
    _shading(p, "1F3E6E")
    r = p.add_run(f"  SLIDE {num}   |   {title.upper()}")
    r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = WHITE; r.font.name = "Calibri"

def tag(doc, label, color_hex="ED7D31"):
    """Small orange/blue inline tag line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    _shading(p, color_hex)
    r = p.add_run(f"  {label}  ")
    r.bold = True; r.font.size = Pt(9)
    r.font.color.rgb = WHITE; r.font.name = "Calibri"

def speaking(doc, text):
    """Main speaking text — slightly indented."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(0.5)
    p.paragraph_format.space_before  = Pt(3)
    p.paragraph_format.space_after   = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(text)
    r.font.size = Pt(11); r.font.name = "Calibri"

def bullet_item(doc, text, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(10.5); r.font.name = "Calibri"
    if color: r.font.color.rgb = color

def qa_item(doc, question, answer):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run("Q: ")
    r1.bold = True; r1.font.color.rgb = NAVY; r1.font.size = Pt(10.5)
    r2 = p.add_run(question)
    r2.bold = True; r2.font.size = Pt(10.5); r2.font.name = "Calibri"
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.8)
    p2.paragraph_format.space_after = Pt(5)
    r3 = p2.add_run("A: ")
    r3.bold = True; r3.font.color.rgb = ORANGE; r3.font.size = Pt(10.5)
    r4 = p2.add_run(answer)
    r4.font.size = Pt(10.5); r4.font.name = "Calibri"

def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(f"📌  Note: {text}")
    r.italic = True; r.font.size = Pt(9.5)
    r.font.color.rgb = GRAY; r.font.name = "Calibri"

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
r = p.add_run("MSc THESIS DEFENSE")
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = ORANGE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
r = p.add_run("PRESENTATION SCRIPT & SPEAKER NOTES")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY

hr(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
r = p.add_run(
    "BEHAVIORAL AND ACADEMIC INDICATOR-BASED DEPRESSION\n"
    "SCREENING AMONG HIGHER EDUCATION STUDENTS:\n"
    "A COMPARATIVE MACHINE LEARNING STUDY"
)
r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
for label, value in [
    ("Presented by:", "Md. Murad Hossain"),
    ("Supervised by:", "Dr. Mrinal Kanti Baowaly"),
    ("Department:", "Computer Science & Engineering"),
    ("Institution:", "Gopalganj Science and Technology University, Bangladesh"),
    ("Date:", "March 2026"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    r1 = p.add_run(f"{label}  ")
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = GRAY
    r2 = p.add_run(value)
    r2.font.size = Pt(11); r2.font.name = "Calibri"

hr(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Total Slides: 14   |   Estimated Presentation Time: 15–20 minutes")
r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GRAY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# HOW TO USE THIS DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
r = p.add_run("HOW TO USE THIS SCRIPT")
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY

speaking(doc,
    "This document contains your complete presentation script — one section per slide. "
    "Each section includes: (1) what to say while showing the slide, "
    "(2) key points to emphasize, and (3) likely examiner questions with suggested answers. "
    "Read through this document several times before the defense until the content feels natural.")

tag(doc, "IMPORTANT TIPS", "1F3E6E")
for tip in [
    "Speak clearly and at a moderate pace — do not rush through slides.",
    "Make eye contact with the examiner, not the screen.",
    "Slides 10–12 (Results) are the most important — know the numbers by heart.",
    "If asked something outside your knowledge, say: 'That is a great point and an interesting direction for future research.'",
    "Time target: ~1.5 minutes per slide — do not exceed 20 minutes total.",
]:
    bullet_item(doc, tip, NAVY)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 1 — TITLE
# ══════════════════════════════════════════════════════════════════════════════
slide_banner(doc, 1, "Title Slide")
tag(doc, "DURATION: ~30 seconds", "2E74B5")

speaking(doc,
    "Good [morning/afternoon], respected members of the examination committee. "
    "I am Md. Murad Hossain, a postgraduate student in the Department of Computer Science and Engineering "
    "at Gopalganj Science and Technology University. "
    "Today, I am presenting my MSc thesis titled: "
    "'Behavioral and Academic Indicator-Based Depression Screening Among Higher Education Students: "
    "A Comparative Machine Learning Study.' "
    "This research was conducted under the supervision of Dr. Mrinal Kanti Baowaly. "
    "I will walk you through the complete study — from the problem we addressed, "
    "to the methods we used, and the results we achieved.")

hr(doc, "ED7D31")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 2 — OUTLINE
# ══════════════════════════════════════════════════════════════════════════════
slide_banner(doc, 2, "Presentation Outline")
tag(doc, "DURATION: ~30 seconds", "2E74B5")

speaking(doc,
    "Here is a brief outline of what I will cover today. "
    "I will begin with the problem statement and research motivation, "
    "followed by our research objectives. "
    "I will then discuss relevant prior work in the literature, "
    "describe our dataset and features, "
    "and explain the methodology including data preprocessing and classification algorithms. "
    "Finally, I will present the experimental results, "
    "discuss the key findings, and close with the conclusion and future directions. "
    "I welcome your questions at the end.")

hr(doc, "ED7D31")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 3 — PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════════════════════
slide_banner(doc, 3, "Problem Statement & Motivation")
tag(doc, "DURATION: ~1.5 minutes", "2E74B5")

speaking(doc,
    "Depression is a serious and growing global public health crisis. "
    "According to the World Health Organization, approximately 280 million people worldwide "
    "are currently affected by depression — making it the leading contributor to disability-adjusted life years globally. "
    "In higher education, the situation is particularly alarming. "
    "Research across South and Southeast Asian institutions shows depression prevalence rates "
    "between 23 and 41 percent among students — significantly higher than the general population. "
    "This is largely driven by intense academic pressure, financial insecurity, career uncertainty, "
    "and rapid social changes that today's students face.")

speaking(doc,
    "In Bangladesh specifically, over 1.5 million students are enrolled in public and private degree programs. "
    "Yet, systematic mental health screening at the institutional level is virtually non-existent. "
    "Traditional clinical tools such as the PHQ-9 questionnaire and psychiatric consultation "
    "are effective, but they require trained clinicians, are time-intensive, and — critically — "
    "many students avoid them due to social stigma around mental health. "
    "This creates a clear gap: we need an automated, low-cost, non-invasive screening tool "
    "that can identify at-risk students before their condition escalates.")

tag(doc, "KEY POINT TO EMPHASIZE", "ED7D31")
bullet_item(doc, "The problem is real, large-scale, and directly relevant to Bangladeshi universities.")
bullet_item(doc, "Existing solutions are not scalable — machine learning offers a practical alternative.")

tag(doc, "LIKELY EXAMINER QUESTIONS", "C00000")
qa_item(doc,
    "Why is depression particularly severe among students in Bangladesh?",
    "Bangladesh's rapidly changing socioeconomic landscape means students face simultaneous pressures: "
    "highly competitive academic environments, financial burden, urbanization-driven family separation, "
    "and limited institutional mental health support. These converging factors make students uniquely vulnerable.")

hr(doc, "ED7D31")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 4 — RESEARCH OBJECTIVES
# ══════════════════════════════════════════════════════════════════════════════
slide_banner(doc, 4, "Research Objectives")
tag(doc, "DURATION: ~1 minute", "2E74B5")

speaking(doc,
    "Our research had five specific objectives. "
    "First, we aimed to construct a clinically-informed dataset by designing a structured "
    "behavioral and academic survey administered to 539 higher education students, "
    "capturing 23 feature dimensions covering psychological, academic, social, and lifestyle indicators. "
    "Second, we designed and implemented a systematic preprocessing pipeline "
    "to clean, encode, and normalize the raw survey data. "
    "Third, we trained four supervised machine learning classifiers — "
    "Support Vector Machine, Logistic Regression, XGBoost, and Multilayer Perceptron. "
    "Fourth, we rigorously compared all four models using five evaluation metrics "
    "under 10-fold stratified cross-validation to ensure robust and unbiased results. "
    "And fifth, we aimed to identify the best-performing model suitable "
    "for real-world institutional deployment.")

tag(doc, "KEY POINT TO EMPHASIZE", "ED7D31")
bullet_item(doc, "The multi-classifier comparison is a key contribution — we do not just train one model.")
bullet_item(doc, "10-fold stratified CV ensures results are statistically reliable, not just a single lucky split.")

hr(doc, "ED7D31")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 5 — LITERATURE REVIEW
# ══════════════════════════════════════════════════════════════════════════════
slide_banner(doc, 5, "Literature Review")
tag(doc, "DURATION: ~2 minutes", "2E74B5")

speaking(doc,
    "Several researchers have explored machine learning for depression detection in recent years. "
    "Mohd and Yahya in 2018 applied Logistic Regression and Artificial Neural Networks "
    "to student questionnaire data, achieving accuracies in the range of 72 to 78 percent. "
    "Their work established that self-reported behavioral and demographic features "
    "carry significant predictive power — directly motivating our survey-based approach.")

speaking(doc,
    "Deshpande and Rao in 2017 used Support Vector Machine and Naive Bayes classifiers "
    "on emotion-tagged social media content, evaluated using F1-score and precision. "
    "Cacheda et al. in 2019 developed a dual Random Forest model on social network activity streams, "
    "demonstrating that combining population-level and individual-level models improves early detection. "
    "Bae and Lee's 2015 study on Korean adolescent clinical data found that "
    "depression severity was the single strongest predictor of suicide attempts — "
    "a finding that reinforced the clinical importance of accurate depression classification. "
    "And Dipnall et al. in 2016 fused machine learning with traditional statistics "
    "to identify biomarkers of depression from large clinical registries.")

speaking(doc,
    "While these studies establish the feasibility of ML-based depression screening, "
    "a critical research gap remains: no published study has offered a "
    "multi-classifier benchmarking framework built specifically on behavioral survey data "
    "tailored to higher education students in Bangladesh. "
    "Our study addresses this gap directly.")

tag(doc, "LIKELY EXAMINER QUESTIONS", "C00000")
qa_item(doc,
    "How is your work different from previous studies?",
    "Previous studies either used social media data — which raises privacy concerns — "
    "or focused on single classifiers. We contribute a clinically-guided, survey-based dataset "
    "specifically from Bangladeshi students, compare four complementary ML paradigms "
    "under identical rigorous experimental conditions, and achieve significantly higher accuracy "
    "than prior student depression prediction work.")

hr(doc, "ED7D31")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 6 — DATASET & FEATURES
# ══════════════════════════════════════════════════════════════════════════════
slide_banner(doc, 6, "Dataset & Feature Description")
tag(doc, "DURATION: ~1.5 minutes", "2E74B5")

speaking(doc,
    "Our dataset was constructed through a structured questionnaire instrument "
    "developed in consultation with licensed psychiatrists and clinical psychologists. "
    "The survey was administered to students pursuing B.Sc., M.Sc., and honours degree programs "
    "at higher education institutions across Bangladesh, on a voluntary and anonymous basis. "
    "After quality screening to remove incomplete responses, we retained 539 validated student records.")

speaking(doc,
    "Each record contains 23 feature dimensions across five categories: "
    "behavioral and psychological features — such as personality type, solitude comfort, and social norms acceptance; "
    "academic and performance features — including SSC result, HSC result, University CGPA, and extracurricular activities; "
    "digital and lifestyle features — smartphone ownership, daily sleep duration, and work contentment; "
    "family and social environment features — family size, family contentment, and comfortable environment; "
    "and clinical risk indicators — including history of suicide attempts and suicidal thoughts, "
    "along with demographic attributes like age and gender. "
    "The class label is binary: 'Depressed' (40.62% of cases) and 'Not Depressed' (59.38%).")

tag(doc, "KEY POINT TO EMPHASIZE", "ED7D31")
bullet_item(doc, "23 features covering 5 distinct psychological and behavioral domains — this breadth is a strength.")
bullet_item(doc, "Dataset was built with psychiatric consultation — it is clinically grounded, not arbitrary.")
bullet_item(doc, "Class distribution is 40.62% / 59.38% — moderately imbalanced but manageable.")

tag(doc, "LIKELY EXAMINER QUESTIONS", "C00000")
qa_item(doc,
    "Why only 539 samples? Is this sufficient for machine learning?",
    "539 samples is a modest but adequate dataset size for the feature dimensionality involved. "
    "We mitigated small-sample risks by using 10-fold stratified cross-validation rather than a single train-test split, "
    "which provides more statistically stable performance estimates. "
    "Furthermore, all four classifiers showed consistent high accuracy, confirming the dataset is sufficient "
    "for this classification task. Expanding the dataset is listed as a priority in our future work.")
qa_item(doc,
    "How was the depression label assigned?",
    "The depression class label was assigned based on standardized clinical assessment criteria "
    "applied by the consulting psychiatrists as part of the data collection process. "
    "Students were not self-labeling — trained professionals reviewed responses to assign the binary label.")

hr(doc, "ED7D31")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 7 — METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════
slide_banner(doc, 7, "Methodology & System Workflow")
tag(doc, "DURATION: ~2 minutes", "2E74B5")

speaking(doc,
    "Our methodology follows a four-stage end-to-end pipeline as shown in Figure 1 on this slide. "
    "Stage one is Survey Design — where we developed the 23-feature questionnaire with psychiatric guidance. "
    "Stage two is Data Collection — the survey was administered and 539 valid responses were gathered. "
    "Stage three is the Preprocessing Pipeline — this is where we transform raw survey responses "
    "into a clean, numerical, machine-ready format. "
    "Stages four through seven cover model training, cross-validation, evaluation, and generating the risk output.")

speaking(doc,
    "The preprocessing pipeline is critical to achieving high accuracy. "
    "It consists of five steps: "
    "First, Binary Encoding — all Yes/No response columns, of which there are 13, "
    "are converted to integer 1 for affirmative and 0 for negative. "
    "Second, One-Hot Encoding — multi-class categorical variables "
    "like Gender, Personality Type, and Comfortable Environment "
    "are expanded into separate binary indicator columns, "
    "preventing the model from incorrectly treating them as ordinal. "
    "Third, Median Imputation — any missing numerical values are replaced using the column median, "
    "which is resistant to outlier distortion. "
    "Fourth, Z-Score Normalization — all numerical features are standardized "
    "to zero mean and unit variance using StandardScaler, "
    "ensuring no single feature dominates others during training. "
    "After these transformations, our original 23 attributes expand to 26 numeric features for model input.")

tag(doc, "LIKELY EXAMINER QUESTIONS", "C00000")
qa_item(doc,
    "Why did you use median imputation instead of mean imputation?",
    "Survey data often contains bounded categorical-to-numeric conversions that can have extreme values "
    "skewing the mean. Median imputation is a more robust strategy in such cases "
    "because the median is not affected by outliers.")
qa_item(doc,
    "Why normalize features?",
    "Several algorithms we use — particularly SVM and MLP — are sensitive to feature scale. "
    "Without normalization, features with large numerical ranges such as age or CGPA "
    "would dominate distance and gradient calculations.")

hr(doc, "ED7D31")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 8 — ALGORITHMS
# ══════════════════════════════════════════════════════════════════════════════
slide_banner(doc, 8, "Classification Algorithms")
tag(doc, "DURATION: ~2 minutes", "2E74B5")

speaking(doc,
    "We selected four classifiers that span fundamentally different machine learning paradigms, "
    "enabling a meaningful comparative study. "
    "The first is Support Vector Machine with an RBF kernel. "
    "SVM finds the maximum-margin hyperplane that best separates depressed from non-depressed students "
    "in a high-dimensional feature space. "
    "The RBF kernel maps features into an infinite-dimensional space, "
    "allowing non-linear boundaries to be captured with geometric precision. "
    "We configured it with C equal to 10 and gamma set to 'scale'.")

speaking(doc,
    "The second classifier is Logistic Regression. "
    "This is a linear probabilistic model that maps feature combinations through a sigmoid function "
    "to produce class probabilities. "
    "It is our interpretability baseline — simple, fast, and fully explainable. "
    "Configured with L-BFGS solver, C equal to 1.0, and 1000 maximum iterations.")

speaking(doc,
    "Third is XGBoost — an ensemble of gradient-boosted decision trees. "
    "XGBoost builds trees sequentially, each one correcting the errors of the previous, "
    "with built-in L1 and L2 regularization to prevent overfitting. "
    "This makes it particularly effective for datasets where features may interact in complex ways. "
    "We used 200 trees, maximum depth of 6, and learning rate of 0.1.")

speaking(doc,
    "Fourth is the Multilayer Perceptron — a feedforward neural network. "
    "Our MLP has three hidden layers with 128, 64, and 32 neurons respectively, "
    "all using ReLU activation, trained with the Adam optimizer. "
    "It captures hierarchical non-linear representations of the behavioral features, "
    "functioning as our deep learning baseline.")

tag(doc, "LIKELY EXAMINER QUESTIONS", "C00000")
qa_item(doc,
    "Why did you choose these four specific algorithms?",
    "We deliberately selected algorithms representing four complementary paradigms: "
    "kernel-based margin maximization (SVM), linear probabilistic modeling (LR), "
    "ensemble gradient boosting (XGBoost), and neural representation learning (MLP). "
    "This ensures our comparison is comprehensive across the main families of supervised learning, "
    "not just variations of the same approach.")
qa_item(doc,
    "Why not try Random Forest or Naive Bayes?",
    "Random Forest and Naive Bayes are valid alternatives. "
    "However, our literature review showed SVM, XGBoost, LR, and MLP as the most widely validated "
    "classifiers for behavioral and clinical survey data. "
    "Expanding to additional algorithms is included in our future work.")

hr(doc, "ED7D31")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 9 — EXPERIMENTAL SETUP
# ══════════════════════════════════════════════════════════════════════════════
slide_banner(doc, 9, "Experimental Setup & Evaluation Protocol")
tag(doc, "DURATION: ~1 minute", "2E74B5")

speaking(doc,
    "All experiments were implemented in Python 3.9 using scikit-learn 1.2.2 for SVM, "
    "Logistic Regression, and MLP; XGBoost 1.7.5 for the gradient boosting model; "
    "pandas and NumPy for data manipulation; and Matplotlib for visualizations. "
    "The complete dataset of 539 instances was split into 80% training — 431 samples — "
    "and 20% testing — 108 samples — using stratified splitting to preserve class proportions.")

speaking(doc,
    "For robust evaluation, we applied 10-fold stratified cross-validation. "
    "In this protocol, the dataset is divided into 10 equal folds, "
    "each fold serving as the test set once while the remaining nine are used for training. "
    "Stratification ensures that the 40.62% depressed / 59.38% non-depressed class ratio "
    "is preserved in every fold. "
    "Performance is then averaged across all 10 folds, giving us a single reliable estimate "
    "that is far more trustworthy than a single train-test split. "
    "We evaluated each model on five metrics: Accuracy, Average Precision, "
    "Average Recall, Average F1-Score, and Average ROC AUC.")

hr(doc, "ED7D31")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 10 — RESULTS TABLE
# ══════════════════════════════════════════════════════════════════════════════
slide_banner(doc, 10, "Experimental Results — Performance Table")
tag(doc, "DURATION: ~2 minutes  |  MOST IMPORTANT SLIDE", "C00000")

speaking(doc,
    "This is the core results slide. Let me walk you through the numbers carefully. "
    "Looking at the table of performance metrics obtained from 10-fold stratified cross-validation:")

# Results table
headers = ["Metric", "MLP", "Logistic Regression", "XGBoost", "SVM ★ BEST"]
rows = [
    ["Accuracy (%)",       "96.84", "97.77",  "97.59",  "97.96"],
    ["Avg. Precision (%)", "95.47", "97.23",  "96.30",  "97.23"],
    ["Avg. Recall (%)",    "95.37", "97.22",  "96.30",  "97.22"],
    ["Avg. F1-Score (%)",  "95.34", "97.22",  "96.30",  "97.22"],
    ["Avg. ROC AUC (%)",   "99.01", "99.50",  "99.57",  "99.11"],
]

table = doc.add_table(rows=1+len(rows), cols=5)
table.style = "Table Grid"
for ci, h in enumerate(headers):
    cell = table.rows[0].cells[ci]
    _set_cell_bg(cell, "1F3E6E")
    _cell_txt(cell, h, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

for ri, row in enumerate(rows):
    bg = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
    for ci, val in enumerate(row):
        cell = table.rows[ri+1].cells[ci]
        is_svm_best = (ci == 4 and ri < 4)
        is_auc_best = (ci == 3 and ri == 4)
        if is_svm_best or is_auc_best:
            _set_cell_bg(cell, "C6EFCE")
            _cell_txt(cell, val, bold=True, size=9,
                      color=RGBColor(0x37,0x56,0x23),
                      align=WD_ALIGN_PARAGRAPH.CENTER if ci>0 else WD_ALIGN_PARAGRAPH.LEFT)
        else:
            _set_cell_bg(cell, bg)
            _cell_txt(cell, val, size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER if ci>0 else WD_ALIGN_PARAGRAPH.LEFT)

doc.add_paragraph()

speaking(doc,
    "The Support Vector Machine achieved the highest overall accuracy of 97.96 percent, "
    "with 97.23 percent precision, 97.22 percent recall, 97.22 percent F1-score, "
    "and 99.11 percent ROC AUC. "
    "Logistic Regression performed second with 97.77 percent accuracy — "
    "a remarkable result for a linear classifier, confirming the feature space is well-separated. "
    "XGBoost achieved 97.59 percent accuracy and the highest ROC AUC of 99.57 percent, "
    "demonstrating its superior ranking ability. "
    "MLP delivered 96.84 percent accuracy — the lowest but still an excellent result. "
    "Crucially, all four models exceeded 96 percent accuracy and all exceeded 99 percent ROC AUC. "
    "This consistency across very different algorithmic paradigms confirms that our feature set "
    "and preprocessing pipeline are intrinsically powerful for this classification task.")

tag(doc, "NUMBERS TO KNOW BY HEART", "ED7D31")
bullet_item(doc, "SVM: 97.96%  |  LR: 97.77%  |  XGBoost: 97.59%  |  MLP: 96.84%", RED)
bullet_item(doc, "All ROC AUC above 99% — near-perfect discrimination", GREEN)
bullet_item(doc, "Inter-model accuracy range = only 1.12 percentage points", NAVY)

tag(doc, "LIKELY EXAMINER QUESTIONS", "C00000")
qa_item(doc,
    "Why is SVM the best model here?",
    "SVM with the RBF kernel is particularly effective because it finds maximum-margin decision boundaries "
    "in a kernel-induced infinite-dimensional space, which gives it excellent generalization on "
    "moderately-sized datasets like ours. The behavioral survey features contain non-linear interaction effects "
    "that the RBF kernel captures excellently, whereas logistic regression, despite performing well, "
    "is limited to linear boundaries.")
qa_item(doc,
    "The accuracy is very high — could there be overfitting?",
    "This is an important question. We explicitly addressed overfitting through 10-fold stratified cross-validation, "
    "which provides 10 independent test evaluations. If the model were overfitting, "
    "we would see high variance across folds. The consistent, stable results across all 10 folds "
    "confirm that the models are genuinely generalizing, not memorizing.")
qa_item(doc,
    "Why is XGBoost ROC AUC higher than SVM despite lower accuracy?",
    "Accuracy and ROC AUC measure different things. Accuracy counts correct classifications at a fixed threshold. "
    "ROC AUC measures ranking quality across all possible thresholds. XGBoost's ensemble of trees "
    "produces well-calibrated probability scores that rank instances more consistently across thresholds, "
    "giving it the highest AUC, even though SVM classifies more instances correctly at the default threshold.")

hr(doc, "ED7D31")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 11 — RESULTS CHARTS
# ══════════════════════════════════════════════════════════════════════════════
slide_banner(doc, 11, "Results — Visual Analysis")
tag(doc, "DURATION: ~1.5 minutes", "2E74B5")

speaking(doc,
    "This slide presents four visualizations of our results. "
    "The top-left bar chart shows the accuracy comparison — SVM clearly leads at 97.96 percent. "
    "The tight clustering of bars confirms that all four models perform at a similar high level. "
    "The top-right ROC curve plot shows the receiver operating characteristic curves for all four models. "
    "Visually, all curves hug the top-left corner tightly, corresponding to the 99-plus percent AUC values. "
    "This means all models can distinguish depressed from non-depressed students with near-perfect reliability "
    "regardless of the classification threshold chosen.")

speaking(doc,
    "The bottom-left multi-metric chart confirms consistent performance across all five metrics for every model. "
    "There is no significant degradation in recall or F1-score relative to accuracy, "
    "which confirms that no model is simply ignoring one class to get high accuracy. "
    "The bottom-right confusion matrices show the breakdown of true positives, true negatives, "
    "false positives, and false negatives for each model. "
    "The balanced distribution of errors across both classes is important for clinical applications — "
    "you do not want a model that only predicts 'not depressed' to appear accurate on an imbalanced dataset.")

hr(doc, "ED7D31")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 12 — KEY FINDINGS
# ══════════════════════════════════════════════════════════════════════════════
slide_banner(doc, 12, "Key Findings & Discussion")
tag(doc, "DURATION: ~1 minute", "2E74B5")

speaking(doc,
    "To summarize the key findings from our analysis. "
    "First: SVM with 97.96 percent accuracy is the optimal classifier for this task. "
    "Its RBF kernel successfully captures the non-linear behavioral patterns associated with depression. "
    "Second: Logistic Regression at 97.77 percent demonstrates that the preprocessed feature space "
    "is largely linearly separable — a validation of the quality of our preprocessing pipeline. "
    "Third: XGBoost's 99.57 percent ROC AUC shows superior probability calibration, "
    "making it the best choice if ranking-based screening is preferred over hard classification. "
    "Fourth: MLP at 96.84 percent confirms the validity of neural feature representations "
    "for behavioral survey data, even with a relatively small training set. "
    "The overarching conclusion: when behavioral and academic survey data is collected with clinical rigor "
    "and processed through a principled pipeline, machine learning classifiers can achieve "
    "clinical-grade depression screening accuracy.")

hr(doc, "ED7D31")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 13 — CONCLUSION & FUTURE WORK
# ══════════════════════════════════════════════════════════════════════════════
slide_banner(doc, 13, "Conclusion & Future Work")
tag(doc, "DURATION: ~1.5 minutes", "2E74B5")

speaking(doc,
    "In conclusion, this study proposed and validated a machine learning-based depression screening framework "
    "developed from a clinically-informed 23-feature behavioral and academic survey "
    "administered to 539 Bangladeshi higher education students. "
    "We demonstrated that four complementary machine learning classifiers — "
    "SVM, Logistic Regression, XGBoost, and MLP — "
    "all achieve exceptional performance when trained on this feature set, "
    "with accuracies above 96 percent and ROC AUC values exceeding 99 percent. "
    "SVM is recommended as the primary classifier with 97.96 percent accuracy. "
    "The framework is feasible, low-cost, non-invasive, and suitable "
    "for deployment as an automated institutional screening tool.")

speaking(doc,
    "Regarding limitations: the dataset covers 539 students and may not fully represent "
    "all demographic groups or institutional settings. "
    "The binary classification scheme does not differentiate depression severity levels.")

speaking(doc,
    "For future work, we plan five directions: "
    "First, expanding the dataset through multi-institution longitudinal sampling. "
    "Second, extending to multi-class severity classification — mild, moderate, and severe — "
    "aligned with the PHQ-9 clinical scale. "
    "Third, applying transformer architectures and graph neural networks for deeper feature modeling. "
    "Fourth, integrating passive behavioral sensing data such as smartphone usage and wearable sleep metrics. "
    "And fifth, developing and deploying a real-time web application for university counseling services.")

tag(doc, "LIKELY EXAMINER QUESTIONS", "C00000")
qa_item(doc,
    "What is the practical impact of this research?",
    "The practical impact is significant. This framework can be implemented as a web-based survey "
    "that students complete at the start of each semester. The trained SVM model then flags "
    "at-risk students automatically, enabling counselors to prioritize interventions "
    "before depression becomes severe. This is scalable, cost-effective, and respects privacy "
    "since no sensitive biometric or social media data is required.")

hr(doc, "ED7D31")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 14 — THANK YOU / Q&A
# ══════════════════════════════════════════════════════════════════════════════
slide_banner(doc, 14, "Thank You / Q&A")
tag(doc, "DURATION: closing statement ~30 seconds", "2E74B5")

speaking(doc,
    "Thank you very much for your attention. "
    "To summarize the core contribution of this thesis: "
    "we proposed a behavioral and academic indicator-based depression screening framework "
    "for higher education students, achieving 97.96 percent accuracy with SVM, "
    "97.77 percent with Logistic Regression, 97.59 percent with XGBoost, "
    "and 96.84 percent with Multilayer Perceptron, all validated under 10-fold stratified cross-validation "
    "on a 539-instance clinically-labeled dataset with 23 features. "
    "I believe this work makes a meaningful contribution toward scalable, automated "
    "mental health support within Bangladeshi higher education institutions. "
    "I am now happy to take your questions.")

hr(doc, "ED7D31")

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX — ADDITIONAL Q&A PREPARATION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
p = doc.add_paragraph()
r = p.add_run("APPENDIX: ADDITIONAL EXAMINER Q&A PREPARATION")
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY
hr(doc)

tag(doc, "TECHNICAL QUESTIONS", "1F3E6E")
qa_item(doc,
    "What is 10-fold stratified cross-validation and why did you use it?",
    "10-fold CV splits the data into 10 equal parts. Each part is used as a test set once, "
    "while the other 9 parts form the training set. This is repeated 10 times and the average score is reported. "
    "'Stratified' means each fold maintains the same class proportion as the full dataset (40.62% / 59.38%). "
    "We used it because our dataset is relatively small — a single train-test split would give "
    "unreliable estimates, but 10-fold CV provides a stable, well-validated performance measure.")

qa_item(doc,
    "What is ROC AUC and why does it matter?",
    "ROC stands for Receiver Operating Characteristic curve. It plots the True Positive Rate against "
    "the False Positive Rate at every possible classification threshold. "
    "AUC — Area Under the Curve — summarizes the entire curve as a single number between 0 and 1. "
    "AUC = 1.0 is perfect, AUC = 0.5 is random. All our models exceeded 0.99, meaning they can "
    "rank depressed students above non-depressed ones with near-perfect reliability — "
    "this is critical for clinical screening where you need consistent ranking quality, "
    "not just a single threshold decision.")

qa_item(doc,
    "What is the difference between precision and recall, and which is more important here?",
    "Precision is: of all students the model predicted as depressed, what fraction actually are? "
    "Recall (sensitivity) is: of all actually depressed students, what fraction did the model detect? "
    "In a clinical screening context, recall is arguably more important — missing a depressed student "
    "(false negative) can have severe consequences. Our SVM achieved 97.22% recall, "
    "meaning it correctly identifies the vast majority of truly depressed students.")

qa_item(doc,
    "Why did you use a binary label instead of severity levels?",
    "The clinical assessment conducted during data collection produced binary labels — "
    "depressed or not depressed — consistent with standard diagnostic criteria. "
    "Multi-level severity classification requires larger datasets and more granular clinical annotation. "
    "We explicitly identify multi-class severity classification as our primary future work direction.")

qa_item(doc,
    "How would you deploy this system in a real institution?",
    "The system would be deployed as a secure web application where students complete "
    "the 23-question survey at the start of each semester. Responses are preprocessed automatically "
    "and fed to the trained SVM model. Students flagged as high-risk are referred to the counseling service. "
    "The system respects privacy — no personal identifiable information is required beyond the survey responses. "
    "The entire pipeline runs in seconds and requires no clinical staff for the screening stage itself.")

qa_item(doc,
    "Is there any bias in your dataset?",
    "Potential biases include self-selection bias — students who chose to participate "
    "may differ from those who declined. There may also be social desirability bias, "
    "where students underreport sensitive items like suicide attempts. "
    "We addressed this through anonymous, voluntary participation and clinical validation of labels. "
    "Expanding the dataset to broader, more diverse cohorts in future work will further reduce bias.")

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(BASE_DIR, "Thesis_Presentation_Script_Murad_Hossain.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
