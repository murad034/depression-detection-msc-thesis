"""
Generate a fully-formatted IEEE-style thesis DOCX
AN EFFECTIVE METHOD BASED ON HUMAN DAILY ACTIVITIES TO DETERMINE THE RATE OF DEPRESSION
Author: Md. Murad Hossain, Dr. Mrinal Kanti Baowaly
Results: SVM 97.96%, LR 97.77%, XGBoost 97.59%, MLP 96.84%
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

VIS_DIR = os.path.join(os.path.dirname(__file__), "visualizations")

# ─────────────────────────────────────────────────────────────────────────────
# Helper utilities
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=True, bottom=True, left=True, right=True, color="000000", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    sides = {}
    if top:    sides["w:top"] = True
    if bottom: sides["w:bottom"] = True
    if left:   sides["w:left"] = True
    if right:  sides["w:right"] = True
    for side in ["w:top", "w:left", "w:bottom", "w:right"]:
        border = OxmlElement(side)
        if side in sides:
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), sz)
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), color)
        else:
            border.set(qn("w:val"), "none")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, italic=False, size_pt=10, font="Times New Roman",
            color=None, underline=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.name = font
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0,
              line_spacing=None, first_line_indent=None, left_indent=None):
    pf = para.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if line_spacing:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_spacing)
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    if left_indent is not None:
        pf.left_indent = Cm(left_indent)


def add_paragraph(doc, text="", bold=False, italic=False, size_pt=10,
                  align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4,
                  first_line_indent=None, line_spacing=12, color=None,
                  font="Times New Roman"):
    para = doc.add_paragraph()
    para_fmt(para, align=align, space_before=space_before, space_after=space_after,
             line_spacing=line_spacing, first_line_indent=first_line_indent)
    if text:
        add_run(para, text, bold=bold, italic=italic, size_pt=size_pt,
                font=font, color=color)
    return para


def section_heading(doc, number, title, size_pt=10):
    """IEEE-style section heading: II. RELATED WORKS"""
    para = doc.add_paragraph()
    para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=4)
    add_run(para, f"{number}. {title}", bold=True, size_pt=size_pt)
    return para


def subsection_heading(doc, letter, title, size_pt=10):
    """IEEE-style sub-heading: A. Information Gathering"""
    para = doc.add_paragraph()
    para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=2)
    add_run(para, f"{letter}. {title}", bold=True, italic=True, size_pt=size_pt)
    return para


def add_columns_section(doc, num_cols=2):
    """Switch the current section to num_cols columns."""
    section = doc.sections[-1]
    sectPr = section._sectPr
    # Remove any existing cols element
    for child in sectPr.findall(qn("w:cols")):
        sectPr.remove(child)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(num_cols))
    cols.set(qn("w:space"), "360")  # 360 twips = 0.25 inch gap
    cols.set(qn("w:equalWidth"), "1")
    sectPr.append(cols)


def add_continuous_section_break(doc, num_cols=1):
    """Insert a continuous section break and set columns for new section."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(num_cols))
    cols.set(qn("w:space"), "360")
    cols.set(qn("w:equalWidth"), "1")
    sectPr.append(cols)
    # Set type = continuous
    type_el = OxmlElement("w:type")
    type_el.set(qn("w:val"), "continuous")
    sectPr.append(type_el)
    pPr.append(sectPr)
    # Remove visible text
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    return para


def add_image_figure(doc, img_path, caption, width_inches=3.0):
    if os.path.exists(img_path):
        para = doc.add_paragraph()
        para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=2)
        run = para.add_run()
        run.add_picture(img_path, width=Inches(width_inches))

    cap = doc.add_paragraph()
    para_fmt(cap, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6)
    add_run(cap, caption, bold=False, italic=True, size_pt=9)


def add_placeholder_figure(doc, caption, width_inches=3.0):
    """Gray box placeholder for figures without image files."""
    para = doc.add_paragraph()
    para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=2)
    # Use a table as a gray box placeholder
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    cell.width = Inches(width_inches)
    set_cell_bg(cell, "D9D9D9")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Figure Placeholder]")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(80, 80, 80)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(14)

    cap = doc.add_paragraph()
    para_fmt(cap, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=6)
    add_run(cap, caption, bold=False, italic=True, size_pt=9)


def body_para(doc, text, first_indent=0.5, size_pt=10, space_after=4):
    """Standard body paragraph with first-line indent."""
    para = doc.add_paragraph()
    para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=space_after,
             first_line_indent=first_indent, line_spacing=12)
    add_run(para, text, size_pt=size_pt)
    return para


# ─────────────────────────────────────────────────────────────────────────────
# Document setup
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()

# Page setup — A4, narrow margins (IEEE-like)
for section in doc.sections:
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)

# Default paragraph style tweak (no extra spacing)
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(10)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after  = Pt(0)

# ─────────────────────────────────────────────────────────────────────────────
# TITLE BLOCK  (single column)
# ─────────────────────────────────────────────────────────────────────────────

# Conference/journal line
conf = doc.add_paragraph()
para_fmt(conf, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)
add_run(conf, "MSc Thesis — Department of Computer Science & Engineering", 
        italic=True, size_pt=9)

# Title
title_p = doc.add_paragraph()
para_fmt(title_p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=6)
add_run(title_p, "AN EFFECTIVE METHOD BASED ON HUMAN DAILY ACTIVITIES\n"
        "TO DETERMINE THE RATE OF DEPRESSION",
        bold=True, size_pt=14, font="Times New Roman")

# Authors
auth_p = doc.add_paragraph()
para_fmt(auth_p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2)
add_run(auth_p, "Md. Murad Hossain", bold=True, size_pt=10)
add_run(auth_p, ",  ", size_pt=10)
add_run(auth_p, "Dr. Mrinal Kanti Baowaly", bold=True, size_pt=10)

# Affiliation
aff_p = doc.add_paragraph()
para_fmt(aff_p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2)
add_run(aff_p, "Department of Computer Science & Engineering\n"
        "Gopalganj Science and Technology University, Gopalganj, Bangladesh\n"
        "Email: muradmd312@gmail.com, mkbaowaly@gmail.com",
        italic=True, size_pt=9)

# Horizontal rule
hr = doc.add_paragraph()
para_fmt(hr, space_before=4, space_after=4)
pPr = hr._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "6")
bottom.set(qn("w:space"), "1")
bottom.set(qn("w:color"), "000000")
pBdr.append(bottom)
pPr.append(pBdr)

# ─────────────────────────────────────────────────────────────────────────────
# ABSTRACT  (single column, slightly indented)
# ─────────────────────────────────────────────────────────────────────────────

abs_heading = doc.add_paragraph()
para_fmt(abs_heading, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=2)
add_run(abs_heading, "Abstract", bold=True, italic=True, size_pt=10)
add_run(abs_heading,
    " — Stress is one of the biggest problems in today's fast-paced society, and "
    "because people's lifestyles change so quickly, it frequently results in depression. "
    "It can be challenging to diagnose depression, especially when its symptoms coexist "
    "with those of other medical conditions. Investigating personality-based characteristics "
    "that contribute to student depression was the aim of this study. A systematic survey "
    "was used to gather a dataset containing 539 cases and 23 different attributes. By "
    "examining this dataset, we created a system that can detect possible reasons for "
    "depression. To ascertain the prevalence of depression among students, four machine "
    "learning techniques were used: Support Vector Machine (SVM), XGBoost, Logistic "
    "Regression, and Multilayer Perceptron (MLP). Several indicators were used to evaluate "
    "performance, and 10-fold cross-validation was used to guarantee the accuracy of the "
    "findings. Our analysis proved that the created models were accurate and successful in "
    "forecasting depression. Support Vector Machine (SVM) achieved the highest accuracy of "
    "97.96%, followed by Logistic Regression (97.77%), XGBoost (97.59%), and Multilayer "
    "Perceptron (96.84%). Through the use of machine learning and human behavior analysis "
    "methodologies, this study seeks to shed light on the incidence of depression in "
    "student populations.",
    size_pt=10)

kw_p = doc.add_paragraph()
para_fmt(kw_p, space_before=4, space_after=6)
add_run(kw_p, "Keywords: ", bold=True, size_pt=10)
add_run(kw_p,
    "Support Vector Machine Algorithm, XGBoost Algorithm, Multilayer Perceptron (MLP), "
    "Data Mining, Depression, Logistic Regression Algorithm.", italic=True, size_pt=10)

# ─────────────────────────────────────────────────────────────────────────────
# Switch to TWO-COLUMN layout for the body
# ─────────────────────────────────────────────────────────────────────────────
add_continuous_section_break(doc, num_cols=2)
add_columns_section(doc, num_cols=2)

# ─────────────────────────────────────────────────────────────────────────────
# I. INTRODUCTION
# ─────────────────────────────────────────────────────────────────────────────
section_heading(doc, "I", "INTRODUCTION")

body_para(doc,
    "Depression is a condition that can appear in a variety of ways. Depression is a disorder "
    "that is dynamic and can appear for a variety of reasons and under various circumstances. "
    "Brain injuries and extreme mental stress can be caused by depression. According to WHO "
    "estimates, depression impacts about 300 million people globally [2]. To determine the root "
    "cause of depression in people, we apply data mining approaches.")

body_para(doc,
    "Over the past 20 years, data mining techniques have been used more and more in a variety "
    "of fields [3]. One of the most important processes for finding and displaying large volumes "
    "of data is data mining. It offers a path to information discovery as well. Data mining uses "
    "automatic learning algorithms to identify, extract, and learn valuable information from "
    "massive databases [4]. Over the past ten years, medical research has found success with the "
    "use of data mining tools, especially in the fields of biomedicine and neuroscience. Recently, "
    "psychiatry has begun to take advantage of these methods' benefits in order to better "
    "understand how mental illness is inherited [5]. Real-time data was gathered from several "
    "participants in this study, and various data mining techniques were used to determine "
    "activities.")

# ─────────────────────────────────────────────────────────────────────────────
# II. RELATED WORKS
# ─────────────────────────────────────────────────────────────────────────────
section_heading(doc, "II", "RELATED WORKS")

body_para(doc,
    "This study presents a fuzzy logic method for flood detection to promptly notify customers "
    "of the possibility of flooding [6]. According to these paper publishers, the rates of suicide "
    "attempts in the entire group were 9.6%, and the degree of depression was the most powerful "
    "predictor of suicide attempts. Compared to the non-depressive group, the suicide attempt "
    "rates were 2.8 and 5.4 times higher in the depressed and prospective depressed groups, "
    "respectively [7].")

body_para(doc,
    "The analysis of emotions related to depression is done using Natural Language Processing "
    "(NLP). A support vector machine and a Naive Bayes classifier were employed by them in their "
    "attempt to predict class. They used Twitter to gather information. Additionally, the findings "
    "were presented using key categorization criteria, such as confusion matrix, F1-score, and "
    "precision [8]. A paper was produced that analyzed social media data and utilized a random "
    "forest algorithm and two machine learning methods. One algorithm will identify subjects who "
    "are depressed, and another will aid in identifying individuals who are not. Their conclusion "
    "was that the dual model is more beneficial than the single one [9]. Additionally, a researcher "
    "used Twitter data to make predictions and compared the outcome with other publications [10].")

# ─────────────────────────────────────────────────────────────────────────────
# III. METHODOLOGY
# ─────────────────────────────────────────────────────────────────────────────
section_heading(doc, "III", "METHODOLOGY")

body_para(doc,
    "Although depression analysis has not seen many comparative investigations, it is currently "
    "being categorized. Our efforts are guided by these works. This section provides a brief "
    "overview of techniques and the effective tools and algorithms used. A survey has been "
    "conducted to gather information for research purposes from students at various Bangladeshi "
    "universities. We have mentioned 23 attributes that were gathered and used in this study, "
    "and the dataset comprises 539 occurrences. Several approaches have been used in this "
    "investigation, including:")

# Strategies list
for item in ["Information Gathering", "Preprocessing of Data",
             "Algorithmic Implementation", "Simulation Environment"]:
    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=2,
             left_indent=0.7, line_spacing=12)
    add_run(p, f"• {item}", size_pt=10)

# Fig 1 placeholder
add_placeholder_figure(doc, "Fig. 1: Procedure for System Operation", width_inches=2.8)

subsection_heading(doc, "A", "Information Gathering")
body_para(doc,
    "Our questionnaire was created after consulting with many psychiatrists. We gather data "
    "from B.Sc., M.Sc., and Undergraduate students. Upon receiving the responses, we generate "
    "our dataset. In the dataset, there are 539 cases. We enumerated 23 characteristics that "
    "were gathered and applied in this study. In Table I, the features list is displayed.")

# ─── TABLE I ─────────────────────────────────────────────────────────────────
tbl_cap = doc.add_paragraph()
para_fmt(tbl_cap, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=2)
add_run(tbl_cap, "TABLE I: Features List", bold=True, size_pt=9)

features = [
    ("Age", "–", "Min: 17, Max: 33"),
    ("Gender", "Male", "66.50%"),
    ("", "Female", "33.50%"),
    ("Do you accept all social norms?", "Yes", "66.03%"),
    ("", "No", "33.97%"),
    ("What are your thoughts on yourself?", "Extrovert", "57.15%"),
    ("", "Introvert", "42.85%"),
    ("Do you own a smartphone?", "Yes", "93.10%"),
    ("", "No", "6.90%"),
    ("Do you enjoy hanging out with friends?", "Yes", "89.25%"),
    ("", "No", "10.75%"),
    ("Are you content in your current role?", "Yes", "66.02%"),
    ("", "No", "33.98%"),
    ("Are you anxious about your work?", "Yes", "71.23%"),
    ("", "No", "28.77%"),
    ("Do you talk to someone about your problems?", "Yes", "61.05%"),
    ("", "No", "38.95%"),
    ("Do you find solitude comfortable?", "Yes", "61.22%"),
    ("", "No", "38.78%"),
    ("Did you attempt suicide?", "Yes", "14.83%"),
    ("", "No", "85.17%"),
    ("Do you think suicide is the answer?", "Yes", "12.40%"),
    ("", "No", "87.60%"),
    ("Are you content with your family?", "Yes", "87.00%"),
    ("", "No", "13.00%"),
    ("Do you feel like a burden to your family?", "Yes", "38.75%"),
    ("", "No", "61.25%"),
    ("Where are you more at ease?", "Family", "62.30%"),
    ("", "Friends", "29.85%"),
    ("", "Others", "7.85%"),
    ("Family Size", "4 Members", "36.72%"),
    ("", "5 Members", "20.60%"),
    ("", "6 Members", "21.76%"),
    ("", "Other", "20.92%"),
    ("Do you engage in extracurricular activities?", "Yes", "67.52%"),
    ("", "No", "32.48%"),
    ("Do you find your education system challenging?", "Yes", "67.30%"),
    ("", "No", "32.70%"),
    ("Daily Sleep Duration", "6 Hours", "18.92%"),
    ("", "7 Hours", "27.80%"),
    ("", "8 Hours", "25.90%"),
    ("", "Other", "27.38%"),
    ("SSC Result", "5.00", "46.55%"),
    ("", "4.50", "7.40%"),
    ("", "4.00", "5.00%"),
    ("", "Other", "41.05%"),
    ("HSC Result", "5.00", "21.51%"),
    ("", "4.50", "8.52%"),
    ("", "4.00", "12.42%"),
    ("", "Other", "57.55%"),
    ("University CGPA", "4.00", "2.21%"),
    ("", "3.25", "2.23%"),
    ("", "3.00", "12.83%"),
    ("", "Other", "82.73%"),
    ("Do you suffer from depression? (Class Label)", "Yes", "40.62%"),
    ("", "No", "59.38%"),
]

t1 = doc.add_table(rows=1 + len(features), cols=3)
t1.style = "Table Grid"
t1.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr_row = t1.rows[0]
headers = ["Feature", "Category", "Distribution"]
for i, h in enumerate(headers):
    cell = hdr_row.cells[i]
    set_cell_bg(cell, "1F3864")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.name = "Times New Roman"

# Data rows
for row_i, (feat, cat, dist) in enumerate(features):
    row = t1.rows[row_i + 1]
    bg = "EAECF0" if row_i % 2 == 0 else "FFFFFF"
    texts = [feat, cat, dist]
    for col_i, txt in enumerate(texts):
        cell = row.cells[col_i]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(txt)
        run.font.size = Pt(8)
        run.font.name = "Times New Roman"

# Column widths
col_widths = [Cm(4.5), Cm(2.2), Cm(1.8)]
for col_i, width in enumerate(col_widths):
    for cell in t1.columns[col_i].cells:
        cell.width = width

subsection_heading(doc, "B", "Preprocessing of Data")
body_para(doc,
    "We have addressed and cleaned the missing data using a variety of methods, including "
    "addressing data discrepancies, smoothing noisy data, and replacing missing values or rows. "
    "The following preprocessing steps were applied:")

prep_steps = [
    ("1. Binary Encoding:", "All Yes/No categorical columns (13 features) were converted to binary values (Yes=1, No=0)."),
    ("2. One-Hot Encoding:", "Multi-category columns (Gender, Personality Type, Comfortable Environment) were expanded into separate binary columns."),
    ("3. Missing Value Imputation:", "Missing numeric values were replaced using median imputation, which is robust to outliers."),
    ("4. Feature Scaling:", "All features were standardized using StandardScaler (zero mean, unit variance)."),
    ("5. Final Feature Space:", "After encoding, the 23 original attributes were transformed into 26 features for model input."),
]

for step_bold, step_text in prep_steps:
    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=2,
             left_indent=0.5, line_spacing=12)
    add_run(p, step_bold + " ", bold=True, size_pt=9)
    add_run(p, step_text, size_pt=9)

add_placeholder_figure(doc, "Fig. 2: Data Preparation Procedure", width_inches=2.8)

subsection_heading(doc, "C", "Algorithmic Implementation")
body_para(doc,
    "For classification, we employed four different methods: Multilayer Perceptron (MLP), "
    "Logistic Regression, XGBoost, and Support Vector Machine (SVM).")

algo_items = [
    ("1) Multilayer Perceptron (MLP):",
     "An example of a multilayer perceptron (MLP) is the Feed-forward Artificial Neural Network "
     "(ANN) class [11]. An input layer, an output layer, and hidden layers make up an MLP. In this "
     "implementation, the architecture is: Input (26 features) → Hidden Layer 1 (128 neurons) → "
     "Hidden Layer 2 (64 neurons) → Hidden Layer 3 (32 neurons) → Output (2 classes). ReLU "
     "activation with Adam optimizer (max_iter=500)."),
    ("2) XGBoost Algorithm:",
     "An approach for supervised learning called XGBoost predicts a target variable by combining "
     "the estimates of simpler models. XGBoost makes use of distributed and parallel computing, "
     "together with a novel tree learning technique [14]. Parameters: 200 estimators, max_depth=6, "
     "learning_rate=0.1."),
    ("3) Logistic Regression Algorithm:",
     "Based on a collection of input factors, logistic regression predicts the likelihood of a "
     "binary result using a mathematical technique. A linear combination of input features is "
     "converted into a probability value between 0 and 1 using the logistic function [1]. "
     "Parameters: max_iter=1000, C=1.0."),
    ("4) Support Vector Machine (SVM):",
     "A supervised machine learning approach that classifies data using statistical learning "
     "theory. SVMs maximize the margin between classes to identify a hyperplane that divides them "
     "in a high-dimensional space [16]. Parameters: kernel=RBF, C=10, gamma=scale."),
]

for bold_part, text_part in algo_items:
    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=2, space_after=2,
             line_spacing=12)
    add_run(p, bold_part + " ", bold=True, size_pt=10)
    add_run(p, text_part, size_pt=10)

subsection_heading(doc, "D", "Simulation Environment")
for line in [
    "Programming Language: Python 3.9",
    "Libraries: scikit-learn 1.2.2, xgboost 1.7.5, pandas 1.5.3, numpy 1.24.3",
    "Validation: 10-Fold Stratified Cross-Validation",
    "Train/Test Split: 80% training (431 samples), 20% testing (108 samples)",
]:
    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=2,
             left_indent=0.5, line_spacing=12)
    add_run(p, f"• {line}", size_pt=10)

# ─────────────────────────────────────────────────────────────────────────────
# IV. EXPERIMENTED OUTCOMES AND ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────
section_heading(doc, "IV", "EXPERIMENTED OUTCOMES AND ANALYSIS")

subsection_heading(doc, "A", "Accuracy")
body_para(doc, "The accuracy is defined as:", space_after=2)

# Formula
form = doc.add_paragraph()
para_fmt(form, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=4, line_spacing=14)
add_run(form, "Accuracy = (TP + TN) / (TP + TN + FP + FN)", italic=True, size_pt=10)

subsection_heading(doc, "B", "Recall / True Positive Rate (Sensitivity)")
body_para(doc,
    "A test's true positive rate (TPR) is the likelihood that individuals with the condition will "
    "yield true-positive results [14].", space_after=2)
form = doc.add_paragraph()
para_fmt(form, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=4, line_spacing=14)
add_run(form, "Recall = TP / (TP + FN)", italic=True, size_pt=10)

subsection_heading(doc, "C", "Precision")
form = doc.add_paragraph()
para_fmt(form, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=4, line_spacing=14)
add_run(form, "Precision = TP / (TP + FP)", italic=True, size_pt=10)

subsection_heading(doc, "D", "F1-Score")
body_para(doc, "The F1-Score (F-Measure) is defined as [15]:", space_after=2)
form = doc.add_paragraph()
para_fmt(form, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=4, line_spacing=14)
add_run(form, "F1 = 2 × (Precision × Recall) / (Precision + Recall)", italic=True, size_pt=10)

subsection_heading(doc, "E", "Discussions")

body_para(doc,
    "The results of all four methods are displayed in Table II and Figure 3. We employed "
    "Multilayer Perceptron (MLP), Logistic Regression, XGBoost, and Support Vector Machine "
    "Algorithm for the binary classification task: predicting whether a student suffers from "
    "depression. Accuracy, Precision, Recall, F1-Score, and ROC AUC were obtained using "
    "10-fold stratified cross-validation.")

# ─── TABLE II ────────────────────────────────────────────────────────────────
tbl2_cap = doc.add_paragraph()
para_fmt(tbl2_cap, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=2)
add_run(tbl2_cap, "TABLE II: Experimented Outcome (10-Fold Stratified Cross-Validation)",
        bold=True, size_pt=9)

metrics = [
    ("Accuracy (%)",           "96.84", "97.77", "97.59", "97.96"),
    ("Average Precision (%)",  "95.47", "97.23", "96.30", "97.23"),
    ("Average Recall (%)",     "95.37", "97.22", "96.30", "97.22"),
    ("Average F1-Score (%)",   "95.34", "97.22", "96.30", "97.22"),
    ("Average ROC AUC (%)",    "99.01", "99.50", "99.57", "99.11"),
]

t2 = doc.add_table(rows=1 + len(metrics), cols=5)
t2.style = "Table Grid"
t2.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
hdr2 = t2.rows[0]
for ci, h in enumerate(["Metric", "MLP", "Logistic\nRegression", "XGBoost", "SVM"]):
    cell = hdr2.cells[ci]
    set_cell_bg(cell, "1F3864")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.name = "Times New Roman"

# Best values bold-highlighted
best_cols = [4, 4, 4, 4, 3]  # column indices (1-based within data cols) of best value per row

for ri, row_data in enumerate(metrics):
    row = t2.rows[ri + 1]
    bg = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
    # Find best value column
    vals = [float(row_data[i]) for i in range(1, 5)]
    best_ci = vals.index(max(vals)) + 1  # 1-indexed offset into row_data

    for ci, txt in enumerate(row_data):
        cell = row.cells[ci]
        set_cell_bg(cell, "D6E4F0" if ci == best_ci else bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(txt)
        run.font.size = Pt(8.5)
        run.font.name = "Times New Roman"
        if ci == best_ci:
            run.bold = True

# Column widths
t2_widths = [Cm(3.5), Cm(1.5), Cm(1.8), Cm(1.6), Cm(1.5)]
for ci, w in enumerate(t2_widths):
    for cell in t2.columns[ci].cells:
        cell.width = w

doc.add_paragraph()

body_para(doc,
    "The results obtained by applying the Multilayer Perceptron (MLP) algorithm achieved "
    "96.84% Accuracy, 95.47% Precision, 95.37% Recall, 95.34% F1-Score, and 99.01% ROC AUC. "
    "Through the use of the Logistic Regression technique, we obtained 97.77% Accuracy, "
    "97.23% Precision, 97.22% Recall, 97.22% F1-Score, and 99.50% ROC AUC. Using the "
    "XGBoost method, we obtained 97.59% Accuracy, 96.30% Precision, 96.30% Recall, 96.30% "
    "F1-Score, and 99.57% ROC AUC. These results were obtained using the Support Vector "
    "Machine algorithm: 97.96% Accuracy, 97.23% Precision, 97.22% Recall, 97.22% F1-Score, "
    "and 99.11% ROC AUC.")

# Fig 3 — real bar chart
bar_img = os.path.join(VIS_DIR, "model_comparison.png")
add_image_figure(doc, bar_img,
    "Fig. 3: Accuracy Comparison — MLP, Logistic Regression, XGBoost, SVM",
    width_inches=2.8)

# Confusion matrices
cm_img = os.path.join(VIS_DIR, "confusion_matrices.png")
add_image_figure(doc, cm_img,
    "Fig. 4: Confusion Matrices for all four models",
    width_inches=2.8)

# ROC curves
roc_img = os.path.join(VIS_DIR, "roc_curves.png")
add_image_figure(doc, roc_img,
    "Fig. 5: ROC Curves for all four models",
    width_inches=2.8)

body_para(doc,
    "The bar chart (Figure 3) illustrates the accuracy comparison across all four algorithms. "
    "SVM achieves the highest accuracy at 97.96%, closely followed by Logistic Regression at "
    "97.77%, XGBoost at 97.59%, and MLP at 96.84%. The remarkably high ROC AUC scores (all "
    "above 99%) confirm that all four models are capable of near-perfect discrimination between "
    "depressed and non-depressed students across all classification thresholds. These results "
    "demonstrate that the combination of properly collected behavioral survey data with systematic "
    "preprocessing and optimized machine learning algorithms is highly effective for student "
    "depression detection.")

# ─────────────────────────────────────────────────────────────────────────────
# V. CONCLUSION
# ─────────────────────────────────────────────────────────────────────────────
section_heading(doc, "V", "CONCLUSION")

body_para(doc,
    "By examining information obtained from structured surveys, this study sought to determine "
    "the causes of depression in college students. We created predictive models that successfully "
    "categorized people according to depressive symptoms and associated behaviors using machine "
    "learning techniques. Four methods were used and assessed: Support Vector Machine (SVM), "
    "XGBoost, Logistic Regression, and Multilayer Perceptron (MLP).")

body_para(doc,
    "With 97.96% classification accuracy, the SVM algorithm outperformed the others, closely "
    "followed by Logistic Regression (97.77%), XGBoost (97.59%), and MLP (96.84%). All four "
    "models achieved ROC AUC scores above 99%, confirming their strong discriminative capability. "
    "These findings show that machine learning methods have significant potential for use in the "
    "early identification and assessment of depression, which might be crucial for prompt "
    "intervention and mental health assistance.")

body_para(doc,
    "Nevertheless, there were restrictions when gathering the data. The lengthy questionnaire was "
    "occasionally resisted by participants, which might have affected the caliber of their answers. "
    "Additionally, the sample size was somewhat limited, indicating that future research will "
    "require larger datasets.")

body_para(doc, "In our upcoming efforts, we want to:", space_after=2)
future_items = [
    "Increase the size of the dataset by adding more people with a range of backgrounds.",
    "Improve prediction performance by implementing deep learning techniques.",
    "Sort depression severity into three categories: severe, medium, and low.",
    "Examine other algorithms, such as Random Forest and Naive Bayes, along with sophisticated pre-processing methods.",
    "Deploy the system as a real-time web application for university counseling services.",
]
for item in future_items:
    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=2,
             left_indent=0.5, line_spacing=12)
    add_run(p, f"• {item}", size_pt=10)

# ─────────────────────────────────────────────────────────────────────────────
# REFERENCES
# ─────────────────────────────────────────────────────────────────────────────
section_heading(doc, "", "REFERENCES")

references = [
    "[1] N. Mohd and Y. Yahya, \"A Data Mining Approach for Prediction of Students' Depression Using Logistic Regression And Artificial Neural Network\", IMCOM'18, ACM, Article 52, 2018. DOI: https://doi.org/10.1145/3164541.3164604",
    "[2] \"Depression\", Who.int, 2020. Available: https://www.who.int/en/news-room/fact-sheets/detail/depression. [Accessed: 06-Nov-2020].",
    "[3] J. Dipnall et al., \"Fusing Data Mining, Machine Learning and Traditional Statistics to Detect Biomarkers Associated with Depression\", PLoS ONE, vol. 11, no. 2, p. e0148195, 2016.",
    "[4] M. Piroomnia et al., \"Data mining approaches for genome-wide association of mood disorders\", Psychiatric Genetics, vol. 22, no. 2, pp. 55-61, 2012.",
    "[5] H. Ni et al., \"Data mining-based study on sub-mentally healthy state among residents in eight provinces and cities in China\", Journal of Traditional Chinese Medicine, vol. 34, no. 4, pp. 511-517, 2014.",
    "[6] Z. Idris and M. Nazir, \"Prediction of Flood Detection System: Fuzzy Logic Approach\", Int'l Journal of Enhanced Research in Science Technology & Engineering, vol. 3, no. 1, 2014.",
    "[7] S. Man Bae and S. A Lee, \"Prediction by data mining, of suicide attempts in Korean adolescents: a national study\", Dovepress, 2015.",
    "[8] M. Deshpande and V. Rao, \"Depression detection using emotion artificial intelligence\", ICISS, 2017. DOI: 10.1109/ISS1.2017.8389299",
    "[9] F. Cacheda, D. Fernandez and F. J. Novoa, \"Early Detection of Depression: Social Network Analysis and Random Forest Techniques\", JMIR, vol. 21, 2019.",
    "[10] K. Shrestha, \"Machine Learning for Depression Diagnosis using Twitter data\", Int'l Journal of Computer Engineering in Research Trends, vol. 5, no. 2, 2018.",
    "[11] \"Multilayer perceptron\", En.wikipedia.org, 2020. Available: https://en.wikipedia.org/wiki/Multilayer_perceptron. [Accessed: 06-Nov-2020].",
    "[12] F. Jimenez and C. Martinez, \"Multi-Objective Evolutionary Rule-Based Classification with Categorical Data\", Entropy, 2018.",
    "[13] E. A. De Melo Gomes Soares et al., \"Analysis of the Fuzzy Unordered Rule Induction Algorithm as a Method for Classification\", V CBSF, pp. 17-28, 2018.",
    "[14] \"XGBoost Algorithm: Long May She Reign!\", Medium, 2020. Available: https://towardsdatascience.com/xgboost-algorithm. [Accessed: 06-Nov-2020].",
    "[15] \"Accuracy, Recall, Precision, F-Score & Specificity, which to optimize on?\", Medium, 2020. Available: https://towardsdatascience.com/accuracy-recall-precision. [Accessed: 06-Nov-2020].",
    "[16] B. E. Boser, I. M. Guyon, & V. Vapnik, \"A training algorithm for optimal margin classifiers\", Proc. 5th Annual Workshop on Computational Learning Theory, pp. 144-152, 1992.",
]

for ref in references:
    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=3,
             left_indent=0.5, first_line_indent=-0.5, line_spacing=11)
    add_run(p, ref, size_pt=8.5)

# ─────────────────────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__),
    "AN_EFFECTIVE_METHOD_DEPRESSION_Murad_Hossain.docx")
doc.save(output_path)
print(f"✅ DOCX saved: {output_path}")
